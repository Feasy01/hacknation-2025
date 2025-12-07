from datetime import date, datetime
import os
from io import BytesIO
from pathlib import Path
from typing import Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_karta_wypadku_document(accident_description: str = "") -> Document:
    """Creates a 'Karta Wypadku' document in memory.
    Args:
        accident_description (str): Description of the accident to include in the document.
    Returns:
        Document: The generated document object.
    """
    doc = Document()
    
    # Styl czcionki
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Nagłówek
    p = doc.add_paragraph('ZAKŁAD UBEZPIECZEŃ SPOŁECZNYCH')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    head = doc.add_heading('KARTA WYPADKU Nr ............ / ............', level=1)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # SEKCJA I
    doc.add_heading('I. DANE IDENTYFIKACYJNE PŁATNIKA SKŁADEK', level=2)
    
    doc.add_paragraph('1. Imię i nazwisko lub nazwa: ' + ('.' * 70))
    
    doc.add_paragraph('2. Adres siedziby: ' + ('.' * 70))
    
    doc.add_paragraph('3. NIP: ' + ('.' * 40) + ' REGON: ' + ('.' * 40))
    
    doc.add_paragraph('Dokument tożsamości (dowód osobisty lub paszport):')
    doc.add_paragraph('Rodzaj dokumentu: ' + ('.' * 30) + ' Seria: ' + ('.' * 20) + ' Numer: ' + ('.' * 20))
    
    # SEKCJA II
    doc.add_heading('II. DANE IDENTYFIKACYJNE POSZKODOWANEGO', level=2)
    
    doc.add_paragraph('1. Imię i nazwisko poszkodowanego: ' + ('.' * 70))
    
    doc.add_paragraph('2. PESEL: ' + ('.' * 50))
    
    doc.add_paragraph('Dokument tożsamości (dowód osobisty lub paszport):')
    doc.add_paragraph('Rodzaj dokumentu: ' + ('.' * 30) + ' Seria: ' + ('.' * 20) + ' Numer: ' + ('.' * 20))
    
    doc.add_paragraph('3. Data i miejsce urodzenia: ' + ('.' * 70))
    
    doc.add_paragraph('4. Adres zamieszkania: ' + ('.' * 70))
    
    doc.add_paragraph('5. Tytuł ubezpieczenia wypadkowego (wymienić numer pozycji i pełny tytuł ubezpieczenia społecznego, zgodnie z art. 3 ust. 3 ustawy z dnia 30 października 2002r. o ubezpieczeniu społecznym z tytułu wypadków przy pracy i chorób zawodowych Dz. U. z 2019r. poz 1205)')
    doc.add_paragraph('Nr 8 - wykonywanie zwykłych czynności związanych z prowadzeniem działalności pozarolniczej w rozumieniu przepisów o systemie ubezpieczeń społecznych')
    
    # SEKCJA III
    doc.add_heading('III. INFORMACJE O WYPADKU', level=2)
    
    doc.add_paragraph('1. Data zgłoszenia oraz imię i nazwisko osoby zgłaszającej wypadek:')
    doc.add_paragraph('.' * 100)
    
    doc.add_paragraph('2. Informacje dotyczące okoliczności, przyczyn, czas i miejsca wypadku')
    if accident_description:
        doc.add_paragraph(accident_description)
    else:
        for _ in range(6):
            doc.add_paragraph('.' * 100)
        
    doc.add_paragraph('3. Świadkowie wypadku:')
    doc.add_paragraph('1) ...................................................................................................................................')
    doc.add_paragraph('2) ...................................................................................................................................')
    
    p = doc.add_paragraph('4. Wypadek ')
    p.add_run('JEST / NIE JEST').bold = True
    p.add_run(' wypadkiem przy pracy określonym w art.3 ust. 3 pkt 8 ustawy z dnia 30 października 2002r. o ubezpieczeniu społecznym z tytułu wypadków przy pracy i chorób zawodowych Dz. U. z 2019r. poz 1205 (uzasadnić, jeśli zdarzenia nie uznano za wypadek przy pracy)')
    doc.add_paragraph('.' * 100)
    doc.add_paragraph('.' * 100)
    
    doc.add_paragraph('5. Stwierdzono, że wyłączną przyczyną wypadku było udowodnione naruszenie przez poszkodowanego przepisów dotyczących ochrony życia i zdrowia, spowodowane przez niego umyślnie lub w wskutek rażącego niedbalstwa (podać dowody)')
    doc.add_paragraph('Nie stwierdzono')
    
    doc.add_paragraph('6. Stwierdzono, że poszkodowany, będąc w stanie nietrzeźwości lub pod wpływem środków odurzających lub substancji psychoaktywnych, przyczynił się w znacznym stopniu do spowodowania wypadku (podać dowody, a w przypadku odmowy przez poszkodowanego poddania się badaniu na zawartość tych substancji w organizmie - zamieścić informację o tym fakcie)')
    doc.add_paragraph('.' * 100)
    
    # SEKCJA IV
    doc.add_heading('IV. POZOSTAŁE INFORMACJE', level=2)
    
    doc.add_paragraph('1. Poszkodowanego (członka rodziny) zapoznano z treścią karty wypadku i pouczono o prawie zgłaszania uwag i zastrzeżeń do ustaleń zawartych w karcie wypadku')
    doc.add_paragraph('.' * 60)
    doc.add_paragraph('(Data i podpis poszkodowanego)')

    t = '2. Kartę wypadku sporządzono w dniu: '
    today = date.today().strftime('%d.%m.%Y')
    doc.add_paragraph(t + today)

    doc.add_paragraph('a) Zakład Ubezpieczeń Społecznych')
    doc.add_paragraph('Nazwa podmiotu zobowiązanego do przygotowania karty')
    doc.add_paragraph('b) Sporządzający:')
    
    doc.add_paragraph('3. Przeszkody i trudności uniemożliwiające sporządzenie karty w terminie 14 dni:')
    doc.add_paragraph('.' * 100)
    
    doc.add_paragraph('4. Kartę odebrano w dniu: ................................   ......................')
    doc.add_paragraph('(Podpis uprawnionego)').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph('5. Załączniki:')
    doc.add_paragraph('1. .................................................................')
    doc.add_paragraph('2. .................................................................')
    
    return doc

def create_karta_wypadku_bytes(accident_description: str = "") -> Tuple[bytes, str]:
    """Generates a 'Karta Wypadku' document in memory and returns it as bytes.
    Args:
        accident_description (str): Description of the accident to include in the document.
    Returns:
        tuple[bytes, str]: A tuple containing the document bytes and the filename.
    """
    doc = create_karta_wypadku_document(accident_description)
    
    # Save to BytesIO instead of file
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Generate filename
    filename = f'Karta_Wypadku_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    return buffer.getvalue(), filename

def create_karta_wypadku(accident_description: str = "") -> str:
    """Generates a 'Karta Wypadku' document with provided accident description.
    DEPRECATED: This function saves to disk. Use create_karta_wypadku_bytes() instead.
    Args:
        accident_description (str): Description of the accident to include in the document.
    Returns:
        str: File path of the generated document.
    """
    doc = create_karta_wypadku_document(accident_description)
    
    # Create generated directory if it doesn't exist
    generated_dir = Path(__file__).parent / 'generated'
    generated_dir.mkdir(exist_ok=True)
    
    # Generate filename and full path
    filename = f'Karta_Wypadku_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    file_path = generated_dir / filename
    
    doc.save(str(file_path))
    print(f"Plik '{file_path}' został wygenerowany.")
    
    # Return relative path for the API response
    return f'generated/{filename}'

if __name__ == "__main__":
    create_karta_wypadku()